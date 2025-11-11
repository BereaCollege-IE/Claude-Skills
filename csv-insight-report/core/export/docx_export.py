"""
Word (DOCX) Export Module
Exports reports to Microsoft Word format using python-docx.
"""

from typing import Dict, List, Any
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import sys
sys.path.append(str(Path(__file__).parent.parent))
from narrative import render, generate_limitations_section


def create_styled_document() -> Document:
    """
    Create a document with custom styles.

    Returns:
        Document object with styles applied
    """
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Customize heading styles
    styles = doc.styles

    # Heading 1 style
    try:
        heading1 = styles['Heading 1']
        heading1.font.size = Pt(18)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 51, 102)
    except:
        pass

    # Heading 2 style
    try:
        heading2 = styles['Heading 2']
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(0, 102, 204)
    except:
        pass

    return doc


def add_title_page(doc: Document, settings: Dict[str, Any], project_name: str) -> None:
    """
    Add a title page to the document.

    Args:
        doc: Document object
        settings: Settings dict
        project_name: Name of the project
    """
    # Title
    title = doc.add_heading(f"Data Insight Report: {project_name}", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Spacing
    doc.add_paragraph()

    # Metadata
    org = settings.get('organization_name', 'Organization')
    purpose = settings.get('purpose', 'Data Analysis')
    horizon = settings.get('horizon', 'N/A')

    doc.add_paragraph(f"Prepared for: {settings.get('audience', 'Stakeholders')}")
    doc.add_paragraph(f"Organization: {org}")
    doc.add_paragraph(f"Purpose: {purpose}")
    doc.add_paragraph(f"Decision Horizon: {horizon}")

    # Date
    from datetime import datetime
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # Page break
    doc.add_page_break()


def add_insight_section(
    doc: Document,
    insight: Dict[str, Any],
    narrative_text: str,
    chart_specs: List[Any],
    figure_path: str = None
) -> None:
    """
    Add an insight section with narrative and optional figure.

    Args:
        doc: Document object
        insight: Insight dict
        narrative_text: Generated narrative text
        chart_specs: List of recommended chart specs
        figure_path: Optional path to figure image
    """
    # Add heading with insight title
    doc.add_heading(insight['title'], level=2)

    # Add narrative
    para = doc.add_paragraph(narrative_text)
    para.style = 'Normal'

    # Add figure if provided
    if figure_path and Path(figure_path).exists():
        doc.add_paragraph()
        doc.add_picture(figure_path, width=Inches(6))

        # Add caption
        caption = doc.add_paragraph(f"Figure {insight['id']}: {insight['title']}")
        caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        caption_run = caption.runs[0]
        caption_run.font.size = Pt(10)
        caption_run.font.italic = True

    # Add caveats if present
    if insight.get('caveats'):
        doc.add_paragraph()
        caveats_para = doc.add_paragraph()
        caveats_para.add_run("Note: ").bold = True
        caveats_para.add_run(" ".join(insight['caveats']))
        caveats_para.style = 'Normal'

    doc.add_paragraph()  # Spacing


def write_docx(project_dir: str, project_data: Dict[str, Any]) -> str:
    """
    Write report to Word format.

    Args:
        project_dir: Path to project directory
        project_data: Dict containing all project data including:
            - settings: Settings dict
            - insights: Insights dict
            - narratives: Dict mapping insight IDs to narrative text
            - profile: Profile dict
            - project_name: Name of project

    Returns:
        Path to generated DOCX file
    """
    settings = project_data.get('settings', {})
    insights_data = project_data.get('insights', {})
    narratives = project_data.get('narratives', {})
    profile = project_data.get('profile', {})
    project_name = project_data.get('project_name', 'Untitled')

    # Create document
    doc = create_styled_document()

    # Add title page
    add_title_page(doc, settings, project_name)

    # Add Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        f"This report presents {len(insights_data.get('top_insights', []))} key insights "
        f"from the analysis of {project_name}. Each insight is supported by statistical "
        "evidence and accompanied by recommended visualizations."
    )
    doc.add_paragraph()

    # Add main insights
    doc.add_heading('Key Insights', level=1)

    for insight in insights_data.get('top_insights', []):
        insight_id = insight['id']
        narrative_text = narratives.get(insight_id, insight['rationale'])

        # Find figure path
        figure_path = Path(project_dir) / 'figures' / f"{insight_id}.png"
        if not figure_path.exists():
            figure_path = None

        add_insight_section(
            doc,
            insight,
            narrative_text,
            insight.get('suggested_visuals', []),
            str(figure_path) if figure_path else None
        )

    # Add limitations section
    doc.add_heading('Limitations and Assumptions', level=1)
    limitations_text = generate_limitations_section(profile, settings)
    doc.add_paragraph(limitations_text)
    doc.add_paragraph()

    # Add appendix if there are additional insights
    appendix_insights = insights_data.get('appendix_insights', [])
    if appendix_insights:
        doc.add_page_break()
        doc.add_heading('Appendix: Additional Insights', level=1)

        for insight in appendix_insights[:10]:  # Limit to 10 additional
            doc.add_heading(insight['title'], level=3)
            doc.add_paragraph(insight['rationale'])
            doc.add_paragraph()

    # Save document
    output_path = Path(project_dir) / 'exports' / 'report.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc.save(str(output_path))

    return str(output_path)


if __name__ == '__main__':
    # Example usage
    print("DOCX export module ready")
